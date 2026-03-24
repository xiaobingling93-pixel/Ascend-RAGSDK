#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
-------------------------------------------------------------------------
This file is part of the RAGSDK project.
Copyright (c) 2025 Huawei Technologies Co.,Ltd.

RAGSDK is licensed under Mulan PSL v2.
You can use this software according to the terms and conditions of the Mulan PSL v2.
You may obtain a copy of Mulan PSL v2 at:

         http://license.coscl.org.cn/MulanPSL2

THIS SOFTWARE IS PROVIDED ON AN "AS IS" BASIS, WITHOUT WARRANTIES OF ANY KIND,
EITHER EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO NON-INFRINGEMENT,
MERCHANTABILITY OR FIT FOR A PARTICULAR PURPOSE.
See the Mulan PSL v2 for more details.
-------------------------------------------------------------------------
"""

import os
import unittest

from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from mx_rag.document.loader import DocxLoader


class DocxLoaderTestCase(unittest.TestCase):
    current_dir = os.path.dirname(os.path.realpath(__file__))
    data_dir = os.path.realpath(os.path.join(current_dir, "../../../data"))

    def test_load(self):
        loader = DocxLoader(os.path.join(self.data_dir, "demo.docx"))
        d = loader.load()
        self.assertEqual(1, len(d))

    def test_load_with_image(self):
        loader = DocxLoader(os.path.join(self.data_dir, "demo.docx"))
        d = loader.load()
        self.assertEqual(1, len(d))

    def test_load_and_split(self):
        loader = DocxLoader(os.path.join(self.data_dir, "demo.docx"))
        res = loader.load_and_split(RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100))
        self.assertEqual(1, len(res))

    def test_title(self):
        loader = DocxLoader(os.path.join(self.data_dir, "title.docx"))
        res = loader.load_and_split(RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100))
        self.assertEqual(1, len(res))

    def test_link(self):
        loader = DocxLoader(os.path.join(self.data_dir, "link.docx"))
        res = loader.load_and_split(RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100))
        self.assertEqual(1, len(res))

    def test_word_number(self):
        document = Document()
        document.add_heading('Document Title', 0)

        word_num = 0
        text = "A plain paragraph having some "
        while word_num <= 500000:
            document.add_paragraph(text)
            word_num += len(text)

        test_file = os.path.join(self.data_dir, "page_number_test.docx")
        document.save(test_file)
        with self.assertRaises(ValueError):
            loader = DocxLoader(test_file)
            loader.load_and_split(RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100))


if __name__ == '__main__':
    unittest.main()
